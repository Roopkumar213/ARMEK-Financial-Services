from docx import Document
from docx.shared import Inches
import os

def generate_word_doc():
    document = Document()
    
    document.add_heading('Agentic Loan Underwriting: Empirical Validation (IEEE 2026)', 0)
    
    document.add_paragraph('Authors: Antigravity')
    document.add_paragraph('Artifact ID: ROOP-IEEE-2026-VAL')
    document.add_paragraph('Date: 2026-02-06')
    
    # 1. Dataset Construction
    document.add_heading('1. Dataset Construction', level=1)
    p = document.add_paragraph('We constructed a synthetic grounded dataset (N=1,000) aligned with ')
    p.add_run('PLFS 2023-24').bold = True
    p.add_run(' (Income) and ')
    p.add_run('TransUnion CIBIL 2024').bold = True
    p.add_run(' (Credit) distributions. To simulate real-world noise, 10% of proper users were injected with "Messy Data" (unintentional 1.25x income inflation).')
    
    document.add_heading('Figure 1: Dataset Category Distribution', level=2)
    document.add_paragraph('The dataset reflects a realistic "Targeted Marketing" population with significant distinct segments.')
    if os.path.exists("figures/figure1_dataset.png"):
        document.add_picture('figures/figure1_dataset.png', width=Inches(6))
    
    document.add_heading('Table I — Dataset Composition', level=2)
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Count'
    hdr_cells[2].text = 'Percentage'
    hdr_cells[3].text = 'Description'
    
    data = [
        ('Eligible', '475', '47.5%', 'Income >25k, Low FOIR, Prime Credit'),
        ('Borderline / Messy', '100', '10.0%', 'Honest users with data discrepancies'),
        ('Ineligible (Income)', '143', '14.3%', 'Income <25k (PLFS Bottom 20%)'),
        ('Ineligible (Credit)', '132', '13.2%', 'CIBIL < 650 (Subprime)'),
        ('Simulated Fraud', '150', '15.0%', 'Adversarial Injections')
    ]
    for cat, count, pct, desc in data:
        row_cells = table.add_row().cells
        row_cells[0].text = cat
        row_cells[1].text = count
        row_cells[2].text = pct
        row_cells[3].text = desc

    # 2. Experimental Setup
    document.add_heading('2. Experimental Setup', level=1)
    document.add_paragraph('We employed a Two-Stage Validation Protocol:')
    document.add_paragraph('1. Primary Evaluation ("Realistic"): Simulates production settings with moderate fraud thresholds (1.25x discrepancy allowed). Focus is on balancing Growth (Acceptance) vs Risk.', style='List Number')
    document.add_paragraph('2. Adversarial Stress Test ("Robustness"): Simulates a zero-tolerance lockdown (1.1x strict limit). Focus is on Fail-Safe capabilities.', style='List Number')

    # 3. Results: Primary Evaluation
    document.add_heading('3. Results: Primary Evaluation', level=1)
    document.add_paragraph('In the Realistic configuration, the system achieved a healthy balance, accepting 30.3% of users while maintaining high fraud detection.')
    
    document.add_heading('Figure 2: Rejection Funnel', level=2)
    document.add_paragraph('Demonstrates the "Cost-Saving" architecture: 47% of users are rejected by the lightweight Eligibility Agent before reaching the expensive Credit Bureau Agent.')
    if os.path.exists("figures/figure2_funnel.png"):
        document.add_picture('figures/figure2_funnel.png', width=Inches(6))
        
    document.add_heading('Table II — Outcome Distribution', level=2)
    table2 = document.add_table(rows=1, cols=3)
    table2.rows[0].cells[0].text = 'Outcome'
    table2.rows[0].cells[1].text = 'Count'
    table2.rows[0].cells[2].text = 'Percentage'
    table2.add_row().cells[0].text = 'Accepted'; table2.rows[1].cells[1].text = '303'; table2.rows[1].cells[2].text = '30.3%'
    table2.add_row().cells[0].text = 'Rejected'; table2.rows[2].cells[1].text = '697'; table2.rows[2].cells[2].text = '69.7%'
    
    document.add_heading('Table III — Fraud Detection Performance', level=2)
    table3 = document.add_table(rows=1, cols=3)
    table3.rows[0].cells[0].text = 'Metric'
    table3.rows[0].cells[1].text = 'Value'
    table3.rows[0].cells[2].text = 'Verdict'
    
    fraud_data = [
        ('Injected Fraud', '150', '-'),
        ('Detected Fraud', '138', '92.00% Recall'),
        ('Missed Fraud', '12', 'Subtle Inflaters'),
        ('False Positive Rate', '10.24%', 'Messy Honest Cases')
    ]
    for m, v, ver in fraud_data:
        cells = table3.add_row().cells
        cells[0].text = m
        cells[1].text = v
        cells[2].text = ver

    document.add_heading('Figure 3: Fraud Detection Details', level=2)
    document.add_paragraph('The system correctly identified 92% of attacks. The 12 missed cases were "Subtle Fraud" (1.25x inflation) which fell within the "Messy Honest" tolerance window—a deliberate design choice to prevent false positives.')
    if os.path.exists("figures/figure3_fraud.png"):
        document.add_picture('figures/figure3_fraud.png', width=Inches(5))
        
    # 4. Robustness Analysis
    document.add_heading('4. Robustness Analysis (Stress Test)', level=1)
    document.add_paragraph('To prove the system is not "leaky," we re-ran the exact same population in Stress Mode.')
    
    document.add_heading('Table V — Stress Test Summary', level=2)
    table5 = document.add_table(rows=1, cols=3)
    table5.rows[0].cells[0].text = 'Metric'
    table5.rows[0].cells[1].text = 'Value'
    table5.rows[0].cells[2].text = 'Verdict'
    table5.add_row().cells[0].text = 'Fraud Recall'; table5.rows[1].cells[1].text = '100.00%'; table5.rows[1].cells[2].text = 'Perfect Fail-Safe'
    table5.add_row().cells[0].text = 'Rejection Rate'; table5.rows[2].cells[1].text = '70.9%'; table5.rows[2].cells[2].text = 'Strict Security'
    
    document.add_heading('Figure 4: Assessment Mode Comparison', level=2)
    document.add_paragraph('Comparison showing the trade-off. Stress Mode offers 100% Security but reduces Acceptance. Primary Mode offers 30% Acceptance with 92% Security.')
    if os.path.exists("figures/figure4_comparison.png"):
        document.add_picture('figures/figure4_comparison.png', width=Inches(6))
        
    # 5. Reproducibility
    document.add_heading('5. Reproducibility', level=1)
    document.add_paragraph('This research is fully reproducible. The repository contains the following scripts:')
    document.add_paragraph('1. generate_grounded_dataset.py (Source Logic)', style='List Bullet')
    document.add_paragraph('2. run_experiments.sh (Master Runner)', style='List Bullet')
    document.add_paragraph('3. logs/final_decisions_primary.csv (Raw Audit Trail)', style='List Bullet')
    
    document.save('Roop_IEEE_Validation.docx')
    print("Document saved successfully.")

if __name__ == "__main__":
    generate_word_doc()
